from io import BytesIO
from docx import Document


def make_a_doc(user_data, poll_data):
    document = Document()

    document.add_heading(f'Результаты пользователя id {user_data[0]}', 0)
    text = f"Результаты опроса:\n" \
           f"Ваш ИНН - {poll_data[1]}\n" \
           f"Ваc зовут - {poll_data[2]}\n" \
           f"Ваш возраст - {poll_data[3]} лет\n"
    document.add_paragraph(text)

    # Нам необходимо создать file-like объект куда мы произведем сохранение нашего docx документа после его генерации.
    # Такие объекты создает функция open() когда мы открываем файлы через неё, но в нашем случае мы создаем данный
    # объект напрямую
    doc_byte_file_object = BytesIO()
    # Сохраняем данные в объект BytesIO
    document.save(doc_byte_file_object)
    # Теперь нам необходимо достать байтовое представление данных из этого объекта и они будут готовы к передаче
    doc_byte_data = doc_byte_file_object.getvalue()
    return doc_byte_data

